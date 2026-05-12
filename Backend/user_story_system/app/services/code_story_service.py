import json
import os
import re
import shutil
import tempfile
import uuid
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import UploadFile

from app.agents.code_story_agent import create_code_story_agent
from app.services.history_service import save_history
from app.services.project_story_service import (
    ProjectParseError,
    _build_call_edges,
    _build_graph,
    _compute_depth,
    _decode_text,
    _safe_filename,
    parse_project_symbols,
)

CODE_STORY_DIR = os.path.join("outputs", "code_story")
os.makedirs(CODE_STORY_DIR, exist_ok=True)

SUPPORTED_CODE_EXTENSIONS = {".py", ".js", ".jsx", ".ts", ".tsx", ".vue", ".java", ".txt"}
LANGUAGE_EXT_MAP = {
    "python": ".py",
    "py": ".py",
    "javascript": ".js",
    "js": ".js",
    "typescript": ".ts",
    "ts": ".ts",
    "vue": ".vue",
    "java": ".java",
}


class CodeStoryParseError(ValueError):
    pass


def _guess_ext(language: str = "", filename: str = "") -> str:
    ext = Path(filename or "").suffix.lower()
    if ext in SUPPORTED_CODE_EXTENSIONS:
        return ext
    return LANGUAGE_EXT_MAP.get((language or "").lower(), ".txt")


def _language_from_ext(ext: str, language: str = "") -> str:
    if language:
        return language
    return {
        ".py": "python",
        ".js": "javascript",
        ".jsx": "javascript",
        ".ts": "typescript",
        ".tsx": "typescript",
        ".vue": "vue",
        ".java": "java",
    }.get(ext, "unknown")


def _overall_node(code: str, filename: str, language: str) -> Dict[str, Any]:
    return {
        "id": "CODE-001",
        "name": Path(filename or "code_snippet").stem or "code_snippet",
        "qualified_name": Path(filename or "code_snippet").stem or "code_snippet",
        "type": "code_block",
        "language": language or "unknown",
        "file": filename or "code_snippet.txt",
        "line": 1,
        "signature": "整体代码片段",
        "code_snippet": code[:3000],
        "docstring": "",
        "calls": [],
        "children": [],
        "parent_id": None,
        "module": Path(filename or "code_snippet").stem or "默认模块",
        "depth": 0,
        "split_strategy": "整体分析：未识别到可拆分的函数/方法/类",
    }


def _renumber_nodes(nodes: List[Dict[str, Any]]) -> None:
    old_to_new = {}
    for index, node in enumerate(nodes, start=1):
        old_id = node.get("id")
        new_id = f"CODE-{index:03d}"
        old_to_new[old_id] = new_id
        node["id"] = new_id
        node["split_strategy"] = "函数/方法/类拆分"
    for node in nodes:
        node["children"] = [old_to_new.get(child, child) for child in node.get("children", [])]
        if node.get("parent_id") in old_to_new:
            node["parent_id"] = old_to_new[node["parent_id"]]


def parse_code_units(code: str, language: str = "", filename: str = "") -> Dict[str, Any]:
    if not code or not code.strip():
        raise CodeStoryParseError("代码内容不能为空")

    safe_name = _safe_filename(filename or "code_snippet")
    ext = _guess_ext(language, safe_name)
    if not Path(safe_name).suffix:
        safe_name = f"{safe_name}{ext}"
    actual_language = _language_from_ext(ext, language)

    workdir = tempfile.mkdtemp(prefix="code_story_")
    try:
        source_dir = os.path.join(workdir, "source")
        os.makedirs(source_dir, exist_ok=True)
        code_path = os.path.join(source_dir, safe_name)
        with open(code_path, "w", encoding="utf-8") as f:
            f.write(code)

        try:
            parsed = parse_project_symbols(source_dir)
            nodes = parsed.get("nodes", [])
            if nodes:
                _renumber_nodes(nodes)
                edges = _build_call_edges(nodes)
                _compute_depth(nodes, edges)
                return {
                    "filename": safe_name,
                    "language": actual_language,
                    "split_mode": "function" if len(nodes) > 1 else "single_symbol",
                    "nodes": nodes,
                    "edges": edges,
                    "graph": _build_graph(nodes, edges),
                    "files": parsed.get("files", []),
                }
        except ProjectParseError:
            pass

        nodes = [_overall_node(code, safe_name, actual_language)]
        return {
            "filename": safe_name,
            "language": actual_language,
            "split_mode": "whole",
            "nodes": nodes,
            "edges": [],
            "graph": _build_graph(nodes, []),
            "files": [{"file": safe_name, "language": actual_language, "symbol_count": 1, "module": Path(safe_name).stem}],
        }
    finally:
        shutil.rmtree(workdir, ignore_errors=True)


def _extract_json_object(text: str) -> Dict[str, Any]:
    cleaned = text.strip()
    cleaned = re.sub(r"^```json\s*", "", cleaned)
    cleaned = re.sub(r"^```\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", cleaned)
        if not match:
            raise
        return json.loads(match.group(0))


def _fallback_items(nodes: List[Dict[str, Any]], edges: List[Dict[str, Any]]) -> Dict[str, Any]:
    incoming_map: Dict[str, List[str]] = {}
    outgoing_map: Dict[str, List[str]] = {}
    for edge in edges:
        outgoing_map.setdefault(edge["source"], []).append(edge["target"])
        incoming_map.setdefault(edge["target"], []).append(edge["source"])
    by_id = {node["id"]: node for node in nodes}

    items = []
    for index, node in enumerate(nodes, start=1):
        callers = [by_id[item]["name"] for item in incoming_map.get(node["id"], []) if item in by_id]
        callees = [by_id[item]["name"] for item in outgoing_map.get(node["id"], []) if item in by_id]
        title = f"支持 {node['name']} 对应的系统能力"
        desc = f"系统需要根据 {node.get('signature') or node['name']} 的代码逻辑，完成对应的数据处理、流程控制或业务支撑能力。"
        if callees:
            desc += f"该能力依赖下游代码单元：{', '.join(callees)}。"
        if callers:
            desc += f"该能力被上游代码单元引用：{', '.join(callers)}。"
        items.append({
            "node_id": node["id"],
            "requirement": {
                "title": title,
                "role": "系统用户" if node.get("type") != "code_block" else "开发人员",
                "description": desc,
                "business_rules": [
                    "代码单元应在正常输入下返回符合预期的处理结果。",
                    "代码单元在异常输入或边界条件下应进行可控处理，避免系统流程中断。",
                ],
            },
            "user_story": {
                "id": f"US-{index:03d}",
                "story": f"As a 系统用户, I want 使用 {node['name']} 对应能力完成业务处理, So that 系统流程能够稳定、正确地执行",
                "acceptance_criteria": [
                    f"Given 输入数据符合 {node['name']} 的处理条件，When 执行该代码单元，Then 系统应返回正确结果或完成对应操作。",
                    f"Given 输入数据异常或依赖调用失败，When 执行 {node['name']}，Then 系统应给出可控处理结果并避免未捕获错误。",
                ],
            },
            "technical_reasoning": f"由 {node.get('file')} 第 {node.get('line')} 行的 {node.get('type')} 代码单元逆向得到。",
        })
    return {"summary": "系统已根据代码单元和引用关系生成用户故事。", "items": items}


def _normalize_items(raw: Dict[str, Any], nodes: List[Dict[str, Any]], edges: List[Dict[str, Any]]) -> Dict[str, Any]:
    valid_ids = {node["id"] for node in nodes}
    fallback = _fallback_items(nodes, edges)
    fallback_by_node = {item["node_id"]: item for item in fallback["items"]}

    normalized = []
    for index, item in enumerate(raw.get("items") or [], start=1):
        node_id = item.get("node_id")
        if node_id not in valid_ids:
            continue
        base = fallback_by_node[node_id]
        req = item.get("requirement") or {}
        story = item.get("user_story") or {}
        criteria = story.get("acceptance_criteria") or base["user_story"]["acceptance_criteria"]
        if isinstance(criteria, str):
            criteria = [criteria]
        rules = req.get("business_rules") or base["requirement"]["business_rules"]
        if isinstance(rules, str):
            rules = [rules]
        normalized.append({
            "node_id": node_id,
            "requirement": {
                "title": req.get("title") or base["requirement"]["title"],
                "role": req.get("role") or base["requirement"]["role"],
                "description": req.get("description") or base["requirement"]["description"],
                "business_rules": rules,
            },
            "user_story": {
                "id": story.get("id") or f"US-{index:03d}",
                "story": story.get("story") or base["user_story"]["story"],
                "acceptance_criteria": criteria,
            },
            "technical_reasoning": item.get("technical_reasoning") or base["technical_reasoning"],
        })

    if not normalized:
        return fallback
    return {"summary": raw.get("summary") or fallback["summary"], "items": normalized}


def _attach_relations(nodes: List[Dict[str, Any]], edges: List[Dict[str, Any]], items: List[Dict[str, Any]]) -> None:
    by_id = {node["id"]: node for node in nodes}
    item_by_node = {item["node_id"]: item for item in items}
    for item in items:
        node_id = item["node_id"]
        incoming = []
        outgoing = []
        for edge in edges:
            if edge.get("target") == node_id and edge.get("source") in by_id:
                src = by_id[edge["source"]]
                incoming.append({"node_id": src["id"], "name": src["name"], "label": edge.get("label", "引用")})
            if edge.get("source") == node_id and edge.get("target") in by_id:
                tgt = by_id[edge["target"]]
                outgoing.append({"node_id": tgt["id"], "name": tgt["name"], "label": edge.get("label", "引用")})
        item["incoming"] = incoming
        item["outgoing"] = outgoing
        node = by_id.get(node_id, {})
        item["node"] = {
            "id": node.get("id"),
            "name": node.get("name"),
            "qualified_name": node.get("qualified_name"),
            "type": node.get("type"),
            "language": node.get("language"),
            "file": node.get("file"),
            "line": node.get("line"),
            "signature": node.get("signature"),
            "code_snippet": node.get("code_snippet", ""),
            "split_strategy": node.get("split_strategy", ""),
        }


def _plain(text: Any) -> str:
    return str(text or "")


def _build_markdown(data: Dict[str, Any]) -> str:
    lines = [
        f"# 代码逆向用户故事分析报告：{data.get('filename', '')}",
        "",
        f"## 一、分析概述\n{data.get('summary', '')}",
        "",
        "## 二、统计信息",
        f"- 拆分方式：{data.get('split_mode', '')}",
        f"- 代码单元数：{data.get('statistics', {}).get('unit_count', 0)}",
        f"- 引用关系数：{data.get('statistics', {}).get('edge_count', 0)}",
        f"- 用户故事数：{data.get('statistics', {}).get('story_count', 0)}",
        "",
        "## 三、代码单元到用户故事",
    ]
    for index, item in enumerate(data.get("items", []), start=1):
        node = item.get("node", {})
        req = item.get("requirement", {})
        story = item.get("user_story", {})
        lines.extend([
            f"### {index}. {node.get('name', '')}（{node.get('type', '')}）",
            f"- 位置：{node.get('file', '')}:{node.get('line', '')}",
            f"- 签名：`{node.get('signature', '')}`",
            f"- 需求标题：{req.get('title', '')}",
            f"- 需求描述：{req.get('description', '')}",
            f"- 用户故事：{story.get('story', '')}",
            "",
            "**业务规则/异常场景：**",
        ])
        for rule in req.get("business_rules", []):
            lines.append(f"- {rule}")
        lines.append("")
        lines.append("**验收标准：**")
        for ac in story.get("acceptance_criteria", []):
            lines.append(f"- {ac}")
        lines.extend(["", f"**技术依据：** {item.get('technical_reasoning', '')}", ""])
        if item.get("incoming") or item.get("outgoing"):
            lines.append("**引用关系：**")
            for rel in item.get("incoming", []):
                lines.append(f"- 被 `{rel.get('name')}` 引用")
            for rel in item.get("outgoing", []):
                lines.append(f"- 引用 `{rel.get('name')}`")
            lines.append("")
    lines.append("## 四、函数/代码单元引用关系")
    by_id = {node["id"]: node for node in data.get("nodes", [])}
    if data.get("edges"):
        for edge in data.get("edges", []):
            source = by_id.get(edge.get("source"), {}).get("qualified_name", edge.get("source"))
            target = by_id.get(edge.get("target"), {}).get("qualified_name", edge.get("target"))
            lines.append(f"- {source} -> {target}")
    else:
        lines.append("- 未识别到函数之间的引用关系。")
    return "\n".join(lines)


def _export_markdown(data: Dict[str, Any], path: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        f.write(_build_markdown(data))


def _export_docx(data: Dict[str, Any], path: str) -> None:
    try:
        from docx import Document
    except Exception as exc:
        raise CodeStoryParseError("导出 Word 需要安装 python-docx：pip install python-docx") from exc
    doc = Document()
    doc.add_heading(f"代码逆向用户故事分析报告：{data.get('filename', '')}", 0)
    doc.add_heading("一、分析概述", level=1)
    doc.add_paragraph(data.get("summary", ""))
    doc.add_heading("二、统计信息", level=1)
    stats = data.get("statistics", {})
    for line in [
        f"拆分方式：{data.get('split_mode', '')}",
        f"代码单元数：{stats.get('unit_count', 0)}",
        f"引用关系数：{stats.get('edge_count', 0)}",
        f"用户故事数：{stats.get('story_count', 0)}",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("三、代码单元到用户故事", level=1)
    for index, item in enumerate(data.get("items", []), start=1):
        node = item.get("node", {})
        req = item.get("requirement", {})
        story = item.get("user_story", {})
        doc.add_heading(f"{index}. {node.get('name', '')}（{node.get('type', '')}）", level=2)
        doc.add_paragraph(f"位置：{node.get('file', '')}:{node.get('line', '')}")
        doc.add_paragraph(f"签名：{node.get('signature', '')}")
        doc.add_paragraph(f"需求标题：{req.get('title', '')}")
        doc.add_paragraph(f"需求描述：{req.get('description', '')}")
        doc.add_paragraph(f"用户故事：{story.get('story', '')}")
        doc.add_paragraph("业务规则/异常场景：")
        for rule in req.get("business_rules", []):
            doc.add_paragraph(rule, style="List Bullet")
        doc.add_paragraph("验收标准：")
        for ac in story.get("acceptance_criteria", []):
            doc.add_paragraph(ac, style="List Number")
        doc.add_paragraph(f"技术依据：{item.get('technical_reasoning', '')}")
        if item.get("incoming") or item.get("outgoing"):
            doc.add_paragraph("引用关系：")
            for rel in item.get("incoming", []):
                doc.add_paragraph(f"被 {rel.get('name')} 引用", style="List Bullet")
            for rel in item.get("outgoing", []):
                doc.add_paragraph(f"引用 {rel.get('name')}", style="List Bullet")
    doc.add_heading("四、函数/代码单元引用关系", level=1)
    by_id = {node["id"]: node for node in data.get("nodes", [])}
    if data.get("edges"):
        for edge in data.get("edges", []):
            source = by_id.get(edge.get("source"), {}).get("qualified_name", edge.get("source"))
            target = by_id.get(edge.get("target"), {}).get("qualified_name", edge.get("target"))
            doc.add_paragraph(f"{source} -> {target}", style="List Bullet")
    else:
        doc.add_paragraph("未识别到函数之间的引用关系。")
    doc.save(path)


def _export_pdf(data: Dict[str, Any], path: str) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase import pdfmetrics
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except Exception as exc:
        raise CodeStoryParseError("导出 PDF 需要安装 reportlab：pip install reportlab") from exc

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title = ParagraphStyle("TitleCN", parent=styles["Title"], fontName="STSong-Light", fontSize=18, leading=24)
    heading = ParagraphStyle("HeadingCN", parent=styles["Heading2"], fontName="STSong-Light", fontSize=14, leading=20, spaceBefore=10)
    body = ParagraphStyle("BodyCN", parent=styles["BodyText"], fontName="STSong-Light", fontSize=10.5, leading=17)

    def p(text: Any, style=body):
        return Paragraph(escape(_plain(text)).replace("\n", "<br/>"), style)

    story = [p(f"代码逆向用户故事分析报告：{data.get('filename', '')}", title), Spacer(1, 12)]
    story.extend([p("一、分析概述", heading), p(data.get("summary", "")), Spacer(1, 8)])
    stats = data.get("statistics", {})
    story.extend([
        p("二、统计信息", heading),
        p(f"拆分方式：{data.get('split_mode', '')}"),
        p(f"代码单元数：{stats.get('unit_count', 0)}"),
        p(f"引用关系数：{stats.get('edge_count', 0)}"),
        p(f"用户故事数：{stats.get('story_count', 0)}"),
        Spacer(1, 8),
        p("三、代码单元到用户故事", heading),
    ])
    for index, item in enumerate(data.get("items", []), start=1):
        node = item.get("node", {})
        req = item.get("requirement", {})
        us = item.get("user_story", {})
        story.extend([
            p(f"{index}. {node.get('name', '')}（{node.get('type', '')}）", heading),
            p(f"位置：{node.get('file', '')}:{node.get('line', '')}"),
            p(f"签名：{node.get('signature', '')}"),
            p(f"需求标题：{req.get('title', '')}"),
            p(f"需求描述：{req.get('description', '')}"),
            p(f"用户故事：{us.get('story', '')}"),
            p("业务规则/异常场景："),
        ])
        for rule in req.get("business_rules", []):
            story.append(p(f"• {rule}"))
        story.append(p("验收标准："))
        for ac in us.get("acceptance_criteria", []):
            story.append(p(f"• {ac}"))
        story.append(p(f"技术依据：{item.get('technical_reasoning', '')}"))
        if item.get("incoming") or item.get("outgoing"):
            story.append(p("引用关系："))
            for rel in item.get("incoming", []):
                story.append(p(f"• 被 {rel.get('name')} 引用"))
            for rel in item.get("outgoing", []):
                story.append(p(f"• 引用 {rel.get('name')}"))
        story.append(Spacer(1, 8))
    story.append(p("四、函数/代码单元引用关系", heading))
    by_id = {node["id"]: node for node in data.get("nodes", [])}
    if data.get("edges"):
        for edge in data.get("edges", []):
            source = by_id.get(edge.get("source"), {}).get("qualified_name", edge.get("source"))
            target = by_id.get(edge.get("target"), {}).get("qualified_name", edge.get("target"))
            story.append(p(f"• {source} -> {target}"))
    else:
        story.append(p("未识别到函数之间的引用关系。"))
    SimpleDocTemplate(path, pagesize=A4).build(story)


async def analyze_code_story(code: str, language: str = "", filename: str = "") -> Dict[str, Any]:
    parsed = parse_code_units(code, language, filename)
    nodes = parsed["nodes"]
    edges = parsed["edges"]
    compact_nodes = [
        {
            "id": node["id"],
            "name": node["name"],
            "qualified_name": node.get("qualified_name"),
            "type": node.get("type"),
            "language": node.get("language"),
            "file": node.get("file"),
            "line": node.get("line"),
            "signature": node.get("signature"),
            "code_snippet": node.get("code_snippet", "")[:1600],
            "calls": node.get("calls", [])[:20],
            "depth": node.get("depth", 0),
            "split_strategy": node.get("split_strategy", ""),
        }
        for node in nodes[:80]
    ]
    task = f"""
请直接根据以下代码单元和引用关系，逐个代码单元逆向生成需求和用户故事。

【文件名】
{parsed['filename']}

【拆分方式】
{parsed['split_mode']}

【代码单元】
{json.dumps(compact_nodes, ensure_ascii=False, indent=2)}

【引用关系】
{json.dumps(edges[:200], ensure_ascii=False, indent=2)}
"""
    try:
        agent = create_code_story_agent()
        result = await agent.run(task=task)
        raw = _extract_json_object(result.messages[-1].content)
        story_result = _normalize_items(raw, nodes, edges)
    except Exception:
        story_result = _fallback_items(nodes, edges)

    _attach_relations(nodes, edges, story_result["items"])
    code_story_id = f"CODE-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
    response = {
        "code_story_id": code_story_id,
        "filename": parsed["filename"],
        "language": parsed["language"],
        "split_mode": parsed["split_mode"],
        "summary": story_result["summary"],
        "nodes": nodes,
        "edges": edges,
        "graph": parsed["graph"],
        "items": story_result["items"],
        "statistics": {
            "unit_count": len(nodes),
            "edge_count": len(edges),
            "story_count": len(story_result["items"]),
        },
    }
    with open(os.path.join(CODE_STORY_DIR, f"{code_story_id}.json"), "w", encoding="utf-8") as f:
        json.dump(response, f, ensure_ascii=False, indent=2)

    history_id = save_history("code", {
        **response,
        "title": parsed["filename"],
        "code_preview": code[:500],
    })
    response["history_id"] = history_id
    return response


async def analyze_code_story_file(file: UploadFile, language: str = "") -> Dict[str, Any]:
    filename = _safe_filename(file.filename or "code_file")
    ext = Path(filename).suffix.lower()
    if ext and ext not in SUPPORTED_CODE_EXTENSIONS:
        raise CodeStoryParseError("仅支持上传单个代码文件：.py、.js、.ts、.vue、.java、.txt 等")
    data = await file.read()
    if not data:
        raise CodeStoryParseError("上传代码文件不能为空")
    code = _decode_text(data)
    return await analyze_code_story(code=code, language=language, filename=filename)


def load_code_story_result(code_story_id: str) -> Dict[str, Any]:
    path = os.path.join(CODE_STORY_DIR, f"{code_story_id}.json")
    if not os.path.exists(path):
        raise FileNotFoundError("未找到该代码用户故事分析结果")
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_code_story_item(code_story_id: str, node_id: str) -> Dict[str, Any]:
    data = load_code_story_result(code_story_id)
    item = next((item for item in data.get("items", []) if item.get("node_id") == node_id), None)
    if not item:
        raise FileNotFoundError("未找到对应代码单元的用户故事")
    return {"code_story_id": code_story_id, "item": item}


def export_code_story(code_story_id: str, format: str = "markdown") -> str:
    data = load_code_story_result(code_story_id)
    safe_format = (format or "markdown").lower()
    filename = f"code_user_stories_{code_story_id}"
    if safe_format in {"md", "markdown"}:
        path = os.path.join(CODE_STORY_DIR, f"{filename}.md")
        _export_markdown(data, path)
        return path
    if safe_format == "docx":
        path = os.path.join(CODE_STORY_DIR, f"{filename}.docx")
        _export_docx(data, path)
        return path
    if safe_format == "pdf":
        path = os.path.join(CODE_STORY_DIR, f"{filename}.pdf")
        _export_pdf(data, path)
        return path
    if safe_format == "json":
        path = os.path.join(CODE_STORY_DIR, f"{filename}.json")
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return path
    raise CodeStoryParseError("导出格式仅支持 markdown、docx、pdf、json")
