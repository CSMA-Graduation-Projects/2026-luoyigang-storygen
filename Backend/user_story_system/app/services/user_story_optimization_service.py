import json
import os
import re
import uuid
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List

from fastapi import UploadFile

from app.agents.user_story_optimizer_agent import create_user_story_optimizer_agent
from app.agents.user_story_splitter_agent import create_user_story_splitter_agent

OPTIMIZATION_DIR = os.path.join("outputs", "story_optimization")
os.makedirs(OPTIMIZATION_DIR, exist_ok=True)

SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".md", ".markdown", ".txt"}


class StoryOptimizationParseError(ValueError):
    pass


def _safe_filename(filename: str) -> str:
    name = os.path.basename(filename or "user_story_document")
    return re.sub(r"[^\w.\-\u4e00-\u9fa5]", "_", name)


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk", "gb2312"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise StoryOptimizationParseError("解析 Word 文档需要安装 python-docx：pip install python-docx") from exc

    doc = Document(BytesIO(data))
    parts: List[str] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception:
        try:
            from PyPDF2 import PdfReader
        except Exception as exc:
            raise StoryOptimizationParseError("解析 PDF 文档需要安装 pypdf：pip install pypdf") from exc

    reader = PdfReader(BytesIO(data))
    parts: List[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(f"【第{index}页】\n{text}")
    return "\n\n".join(parts)


async def extract_story_document_text(file: UploadFile) -> Dict[str, str]:
    filename = _safe_filename(file.filename or "user_story_document")
    ext = os.path.splitext(filename)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise StoryOptimizationParseError("仅支持 Word(.docx)、PDF(.pdf)、Markdown(.md/.markdown)、TXT(.txt) 格式")

    data = await file.read()
    if not data:
        raise StoryOptimizationParseError("上传文档内容不能为空")

    if ext == ".docx":
        text = _extract_docx_text(data)
    elif ext == ".pdf":
        text = _extract_pdf_text(data)
    else:
        text = _decode_text(data)

    text = text.strip()
    if not text:
        raise StoryOptimizationParseError("未能从文档中解析出有效文本，请检查文档是否为空或是否为扫描件")

    return {"filename": filename, "extension": ext, "text": text}


def _extract_json_object(text: str) -> Dict[str, Any]:
    cleaned = text.strip()
    cleaned = re.sub(r"^```json\s*", "", cleaned)
    cleaned = re.sub(r"^```\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", cleaned)
        if not match:
            raise
        return json.loads(match.group(0))



def _looks_like_story(line: str) -> bool:
    keywords = (
        "As a", "作为", "我希望", "我想要", "用户故事", "验收标准",
        "I want", "So that", "Given", "When", "Then"
    )
    return any(k.lower() in line.lower() for k in keywords)


def _fallback_split_user_stories(text: str) -> List[Dict[str, str]]:
    """智能体不可用时的兜底拆分，避免接口失败。"""
    normalized = text.replace("\r\n", "\n")

    pattern = re.compile(
        r"(?=(?:^|\n)\s*(?:#{1,6}\s*)?(?:US|Story|用户故事|故事)[-编号\s:：]*\d+)",
        re.IGNORECASE
    )
    chunks = [chunk.strip() for chunk in pattern.split(normalized) if chunk.strip()]

    if len(chunks) <= 1:
        chunks = [chunk.strip() for chunk in re.split(r"\n\s*\n", normalized) if chunk.strip()]

    buffer: List[str] = []
    for chunk in chunks:
        if _looks_like_story(chunk) or len(chunk) > 30:
            buffer.append(chunk)

    if not buffer:
        sentences = [s.strip() for s in re.split(r"[。；;]\s*", normalized) if s.strip()]
        buffer = [s for s in sentences if len(s) > 10]

    stories: List[Dict[str, str]] = []
    for index, chunk in enumerate(buffer, start=1):
        title = ""
        for line in chunk.split("\n"):
            line = line.strip("# -\t ")
            if line:
                title = line[:40]
                break
        stories.append({
            "id": f"OUS-{index:03d}",
            "title": title or f"原始用户故事 {index}",
            "original_story": chunk.strip(),
            "split_reason": "智能体拆分不可用时，系统根据用户故事编号、段落和关键词进行兜底拆分。"
        })

    if not stories:
        stories.append({
            "id": "OUS-001",
            "title": "原始用户故事 1",
            "original_story": normalized[:4000],
            "split_reason": "文档中未检测到明确分隔符，系统将主要内容作为一条用户故事处理。"
        })

    return stories[:50]


def _normalize_split_result(raw: Dict[str, Any], source_text: str) -> List[Dict[str, str]]:
    items = raw.get("stories") or raw.get("items") or []
    if not isinstance(items, list):
        raise ValueError("智能体拆分结果 stories 字段格式错误")

    stories: List[Dict[str, str]] = []
    for index, item in enumerate(items, start=1):
        if isinstance(item, str):
            original_story = item.strip()
            title = original_story.split("\n", 1)[0].strip("# -\t ")[:40]
            split_reason = "智能体根据语义边界识别为独立用户故事。"
        elif isinstance(item, dict):
            original_story = str(item.get("original_story") or item.get("story") or item.get("content") or "").strip()
            title = str(item.get("title") or "").strip()
            split_reason = str(item.get("split_reason") or "智能体根据角色、目标、业务价值和上下文识别为独立用户故事。").strip()
        else:
            continue

        if not original_story:
            continue

        stories.append({
            "id": f"OUS-{index:03d}",
            "title": title or original_story.split("\n", 1)[0].strip("# -\t ")[:40] or f"原始用户故事 {index}",
            "original_story": original_story,
            "split_reason": split_reason
        })

    if not stories:
        return _fallback_split_user_stories(source_text)

    return stories[:50]


async def split_user_stories_by_agent(text: str) -> List[Dict[str, str]]:
    """采用拆分智能体进行语义级用户故事拆分。"""
    agent = create_user_story_splitter_agent()
    task = f"""
请对以下用户故事文档进行智能拆分。

要求：
1. 根据语义完整性拆分，不要简单按换行或编号机械切分。
2. 每条用户故事只包含一个主要角色、一个目标和一个业务价值。
3. 如果一条故事包含多个功能点，请拆成多条。
4. 如果验收标准或补充说明属于某条故事，请合并到该故事 original_story 中。
5. 只输出合法 JSON。

【用户故事文档】
{text[:20000]}
"""
    try:
        result = await agent.run(task=task)
        raw = _extract_json_object(result.messages[-1].content)
        return _normalize_split_result(raw, text)
    except Exception:
        return _fallback_split_user_stories(text)

def _fallback_optimization(original: str, index: int) -> Dict[str, Any]:
    title = original.split("\n", 1)[0].strip("# -\t ")[:40] or f"用户故事 {index}"
    cleaned = re.sub(r"\s+", " ", original).strip()
    want = cleaned[:120] if cleaned else "完成相关业务功能"
    return {
        "title": title,
        "problem_summary": "原用户故事可能存在角色、业务价值或验收条件不够明确的问题。",
        "optimized_story": f"As a 系统用户, I want {want}, So that 能够明确、高效地完成对应业务目标。",
        "acceptance_criteria": [
            "Given 用户具备相应权限，When 用户执行该用户故事对应操作，Then 系统应正确完成业务处理并返回明确结果。",
            "Given 用户输入无效或缺失信息，When 用户提交操作，Then 系统应给出可理解的错误提示并阻止异常数据保存。",
            "Given 操作完成，When 用户查看结果，Then 系统应展示与该操作一致的状态或数据。"
        ],
        "improvement_points": [
            {
                "dimension": "Valuable",
                "before": "业务价值描述不够突出。",
                "after": "补充 So that 价值说明。",
                "reason": "让开发和测试人员理解该故事对用户或业务的意义。"
            },
            {
                "dimension": "Testable",
                "before": "缺少明确验收标准。",
                "after": "补充 Given-When-Then 验收标准。",
                "reason": "便于后续测试用例设计和质量验证。"
            }
        ],
        "quality_dimensions": {
            "Independent": "建议保持该故事可独立开发和验收。",
            "Negotiable": "保留业务目标，避免限定过细实现方案。",
            "Valuable": "突出用户或业务收益。",
            "Estimable": "压缩故事范围，便于估算。",
            "Small": "建议拆分过大的复合功能。",
            "Testable": "增加可验证的验收标准。"
        }
    }


def _normalize_optimization(raw: Dict[str, Any], story: Dict[str, str], index: int) -> Dict[str, Any]:
    criteria = raw.get("acceptance_criteria") or []
    if isinstance(criteria, str):
        criteria = [criteria]

    improvement_points = raw.get("improvement_points") or []
    if isinstance(improvement_points, str):
        improvement_points = [{"dimension": "Overall", "before": "原描述质量不足", "after": improvement_points, "reason": "提升整体表达质量"}]

    normalized_points = []
    for p_idx, point in enumerate(improvement_points, start=1):
        if isinstance(point, dict):
            normalized_points.append({
                "dimension": point.get("dimension") or f"提升点{p_idx}",
                "before": point.get("before") or "优化前表达不够清晰",
                "after": point.get("after") or "优化后表达更明确",
                "reason": point.get("reason") or "提升用户故事质量"
            })
        else:
            normalized_points.append({
                "dimension": f"提升点{p_idx}",
                "before": "优化前表达不够清晰",
                "after": str(point),
                "reason": "提升用户故事质量"
            })

    dims = raw.get("quality_dimensions") or {}
    for key in ["Independent", "Negotiable", "Valuable", "Estimable", "Small", "Testable"]:
        dims.setdefault(key, "已结合该维度进行优化。")

    return {
        "id": story["id"],
        "title": raw.get("title") or story.get("title") or f"用户故事 {index}",
        "original_story": story.get("original_story") or "",
        "problem_summary": raw.get("problem_summary") or "已识别原用户故事在清晰度、价值表达或可测试性方面的改进空间。",
        "optimized_story": raw.get("optimized_story") or story.get("original_story") or "",
        "acceptance_criteria": criteria,
        "improvement_points": normalized_points,
        "quality_dimensions": dims
    }


async def optimize_single_user_story(story: Dict[str, str], index: int) -> Dict[str, Any]:
    agent = create_user_story_optimizer_agent()
    task = f"""
请优化以下第 {index} 条用户故事。要求：保留原始业务意图，按 INVEST 原则提升质量，输出合法 JSON。

【原始用户故事】
{story.get('original_story', '')[:8000]}
"""
    try:
        result = await agent.run(task=task)
        raw = _extract_json_object(result.messages[-1].content)
    except Exception:
        raw = _fallback_optimization(story.get("original_story", ""), index)
    return _normalize_optimization(raw, story, index)


async def analyze_and_optimize_story_document(file: UploadFile) -> Dict[str, Any]:
    parsed = await extract_story_document_text(file)
    original_stories = await split_user_stories_by_agent(parsed["text"])

    optimized_items = []
    for index, story in enumerate(original_stories, start=1):
        optimized_items.append(await optimize_single_user_story(story, index))

    optimization_id = f"OPT-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
    response = {
        "optimization_id": optimization_id,
        "filename": parsed["filename"],
        "extension": parsed["extension"],
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "summary": {
            "story_count": len(optimized_items),
            "improvement_count": sum(len(item.get("improvement_points", [])) for item in optimized_items),
            "criteria_count": sum(len(item.get("acceptance_criteria", [])) for item in optimized_items)
        },
        "items": optimized_items
    }

    with open(os.path.join(OPTIMIZATION_DIR, f"{optimization_id}.json"), "w", encoding="utf-8") as f:
        json.dump(response, f, ensure_ascii=False, indent=2)

    return response


def load_optimization_result(optimization_id: str) -> Dict[str, Any]:
    path = os.path.join(OPTIMIZATION_DIR, f"{optimization_id}.json")
    if not os.path.exists(path):
        raise FileNotFoundError("未找到该用户故事优化结果")
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_optimization_item(optimization_id: str, story_id: str) -> Dict[str, Any]:
    data = load_optimization_result(optimization_id)
    for item in data.get("items", []):
        if item.get("id") == story_id:
            return {"optimization_id": optimization_id, "item": item}
    raise FileNotFoundError("未找到对应用户故事优化项")


def export_optimization_json(optimization_id: str) -> str:
    data = load_optimization_result(optimization_id)
    filename = f"user_story_optimization_{optimization_id}.json"
    path = os.path.join(OPTIMIZATION_DIR, filename)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return path


def export_optimization_markdown(optimization_id: str) -> str:
    data = load_optimization_result(optimization_id)
    filename = f"user_story_optimization_{optimization_id}.md"
    path = os.path.join(OPTIMIZATION_DIR, filename)

    lines = [
        f"# 用户故事优化结果：{data.get('filename', '')}",
        "",
        "## 统计信息",
        f"- 用户故事数量：{data.get('summary', {}).get('story_count', 0)}",
        f"- 提升点数量：{data.get('summary', {}).get('improvement_count', 0)}",
        f"- 验收标准数量：{data.get('summary', {}).get('criteria_count', 0)}",
        ""
    ]

    for item in data.get("items", []):
        lines.extend([
            f"## {item.get('id')} {item.get('title')}",
            "",
            "### 优化前",
            item.get("original_story", ""),
            "",
            "### 主要问题",
            item.get("problem_summary", ""),
            "",
            "### 优化后",
            item.get("optimized_story", ""),
            "",
            "### 验收标准",
        ])
        for criteria in item.get("acceptance_criteria", []):
            lines.append(f"- {criteria}")
        lines.extend(["", "### 提升点"])
        for point in item.get("improvement_points", []):
            lines.append(f"- **{point.get('dimension', '')}**：{point.get('before', '')} → {point.get('after', '')}。原因：{point.get('reason', '')}")
        lines.append("")

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return path


def export_optimization_docx(optimization_id: str) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise StoryOptimizationParseError("导出 Word 文档需要安装 python-docx：pip install python-docx") from exc

    data = load_optimization_result(optimization_id)
    filename = f"user_story_optimization_{optimization_id}.docx"
    path = os.path.join(OPTIMIZATION_DIR, filename)

    doc = Document()
    doc.add_heading(f"用户故事优化结果：{data.get('filename', '')}", level=1)
    summary = data.get("summary", {})
    doc.add_heading("统计信息", level=2)
    doc.add_paragraph(f"用户故事数量：{summary.get('story_count', 0)}")
    doc.add_paragraph(f"提升点数量：{summary.get('improvement_count', 0)}")
    doc.add_paragraph(f"验收标准数量：{summary.get('criteria_count', 0)}")

    for item in data.get("items", []):
        doc.add_heading(f"{item.get('id')} {item.get('title')}", level=2)
        doc.add_heading("优化前", level=3)
        doc.add_paragraph(item.get("original_story", ""))
        doc.add_heading("主要问题", level=3)
        doc.add_paragraph(item.get("problem_summary", ""))
        doc.add_heading("优化后", level=3)
        doc.add_paragraph(item.get("optimized_story", ""))
        doc.add_heading("验收标准", level=3)
        for criteria in item.get("acceptance_criteria", []):
            doc.add_paragraph(criteria, style="List Bullet")
        doc.add_heading("提升点", level=3)
        for point in item.get("improvement_points", []):
            doc.add_paragraph(
                f"{point.get('dimension', '')}：{point.get('before', '')} → {point.get('after', '')}。原因：{point.get('reason', '')}",
                style="List Bullet"
            )

    doc.save(path)
    return path
