import json
import os
import re
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional

from fastapi import UploadFile

from app.agents.document_requirement_agent import create_document_requirement_agent
from app.services.history_service import save_history

DOCUMENT_STORY_DIR = os.path.join("outputs", "document_story")
os.makedirs(DOCUMENT_STORY_DIR, exist_ok=True)

SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".md", ".markdown", ".txt"}


class DocumentParseError(ValueError):
    pass


def _safe_filename(filename: str) -> str:
    name = os.path.basename(filename or "requirement_document")
    return re.sub(r"[^\w.\-\u4e00-\u9fa5]", "_", name)


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk", "gb2312"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
        from io import BytesIO
    except Exception as exc:
        raise DocumentParseError("解析 Word 文档需要安装 python-docx：pip install python-docx") from exc

    doc = Document(BytesIO(data))
    parts: List[str] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        from io import BytesIO
    except Exception:
        try:
            from PyPDF2 import PdfReader
            from io import BytesIO
        except Exception as exc:
            raise DocumentParseError("解析 PDF 文档需要安装 pypdf：pip install pypdf") from exc

    reader = PdfReader(BytesIO(data))
    parts: List[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            parts.append(f"【第{index}页】\n{text}")
    return "\n\n".join(parts)


async def extract_text_from_upload(file: UploadFile) -> Dict[str, str]:
    filename = _safe_filename(file.filename or "requirement_document")
    ext = os.path.splitext(filename)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError("仅支持 Word(.docx)、PDF(.pdf)、Markdown(.md/.markdown)、TXT(.txt) 格式")

    data = await file.read()
    if not data:
        raise DocumentParseError("上传文档内容不能为空")

    if ext == ".docx":
        text = _extract_docx_text(data)
    elif ext == ".pdf":
        text = _extract_pdf_text(data)
    else:
        text = _decode_text(data)

    text = text.strip()
    if not text:
        raise DocumentParseError("未能从文档中解析出有效文本，请检查文档是否为扫描件或空文件")

    return {"filename": filename, "extension": ext, "text": text}


def _extract_json_object(text: str) -> Dict[str, Any]:
    cleaned = text.strip()
    cleaned = re.sub(r"^```json\s*", "", cleaned)
    cleaned = re.sub(r"^```\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", cleaned)
        if not match:
            raise
        return json.loads(match.group(0))


def _fallback_requirement_analysis(text: str) -> Dict[str, Any]:
    """当大模型输出不可解析时，使用规则兜底，保证前端仍可展示。"""
    lines = [line.strip() for line in re.split(r"[\n。；;]", text) if line.strip()]
    keywords = ("用户", "系统", "管理员", "支持", "可以", "能够", "需要", "应", "必须", "提供", "上传", "查询", "导出", "生成")
    nfr_keywords = ("性能", "安全", "响应", "并发", "兼容", "可靠", "可用", "日志", "权限", "加密", "备份")

    requirements = []
    for line in lines:
        if not any(k in line for k in keywords):
            continue
        req_id = f"REQ-{len(requirements) + 1:03d}"
        req_type = "non_functional" if any(k in line for k in nfr_keywords) else "functional"
        story = []
        if req_type == "functional":
            story = [{
                "id": f"US-{len(requirements) + 1:03d}",
                "role": "系统用户",
                "want": line,
                "benefit": "完成对应业务操作并满足业务处理目标",
                "story": f"As a 系统用户, I want {line}, So that 完成对应业务操作并满足业务处理目标",
                "acceptance_criteria": [
                    f"Given 用户具备操作权限，When 用户执行“{line}”，Then 系统应正确完成该功能。",
                    "Given 输入数据不完整或不合法，When 用户提交操作，Then 系统应给出明确提示并阻止错误数据保存。"
                ],
                "notes": "由规则兜底生成，建议人工复核。"
            }]
        requirements.append({
            "id": req_id,
            "title": line[:30],
            "module": "默认模块",
            "type": req_type,
            "priority": "medium",
            "description": line,
            "source_text": line,
            "user_stories": story,
            "non_functional_constraints": [line] if req_type == "non_functional" else []
        })
        if len(requirements) >= 20:
            break

    relations = []
    for index in range(len(requirements) - 1):
        relations.append({
            "source": requirements[index]["id"],
            "target": requirements[index + 1]["id"],
            "type": "related_to",
            "label": "文档相邻需求，存在业务相关性"
        })

    return {
        "summary": "系统根据文档文本识别出若干需求项，并生成结构化用户故事。",
        "modules": ["默认模块"],
        "requirements": requirements,
        "relations": relations
    }


def _normalize_analysis(raw: Dict[str, Any]) -> Dict[str, Any]:
    requirements = raw.get("requirements") or []
    normalized_requirements: List[Dict[str, Any]] = []
    used_ids = set()

    for index, item in enumerate(requirements, start=1):
        req_id = str(item.get("id") or f"REQ-{index:03d}").strip()
        if not req_id.startswith("REQ-") or req_id in used_ids:
            req_id = f"REQ-{index:03d}"
        used_ids.add(req_id)

        req_type = item.get("type") or "functional"
        if req_type not in {"functional", "non_functional"}:
            req_type = "functional"

        stories = item.get("user_stories") or []
        normalized_stories = []
        for story_index, story in enumerate(stories, start=1):
            criteria = story.get("acceptance_criteria") or []
            if isinstance(criteria, str):
                criteria = [criteria]
            normalized_stories.append({
                "id": story.get("id") or f"US-{index:03d}-{story_index}",
                "role": story.get("role") or "系统用户",
                "want": story.get("want") or item.get("title") or item.get("description") or "完成该功能",
                "benefit": story.get("benefit") or "满足业务处理目标",
                "story": story.get("story") or f"As a 系统用户, I want {item.get('title', '完成该功能')}, So that 满足业务处理目标",
                "acceptance_criteria": criteria,
                "notes": story.get("notes") or ""
            })

        normalized_requirements.append({
            "id": req_id,
            "title": item.get("title") or item.get("description") or req_id,
            "module": item.get("module") or "未分组模块",
            "type": req_type,
            "priority": item.get("priority") if item.get("priority") in {"high", "medium", "low"} else "medium",
            "description": item.get("description") or item.get("title") or "",
            "source_text": item.get("source_text") or "",
            "user_stories": normalized_stories,
            "non_functional_constraints": item.get("non_functional_constraints") or []
        })

    valid_ids = {item["id"] for item in normalized_requirements}
    relations = []
    for rel in raw.get("relations") or []:
        source = rel.get("source")
        target = rel.get("target")
        if source in valid_ids and target in valid_ids and source != target:
            relations.append({
                "source": source,
                "target": target,
                "type": rel.get("type") or "related_to",
                "label": rel.get("label") or "相关"
            })

    modules = sorted({item["module"] for item in normalized_requirements})
    return {
        "summary": raw.get("summary") or "已完成需求文档解析与用户故事生成。",
        "modules": raw.get("modules") or modules,
        "requirements": normalized_requirements,
        "relations": relations
    }


def _build_graph(requirements: List[Dict[str, Any]], relations: List[Dict[str, Any]]) -> Dict[str, Any]:
    module_index = {module: idx for idx, module in enumerate(sorted({req["module"] for req in requirements}))}
    nodes = []
    for req in requirements:
        nodes.append({
            "id": req["id"],
            "name": f"{req['id']}\n{req['title']}",
            "category": module_index.get(req["module"], 0),
            "symbolSize": 52 if req["type"] == "functional" else 44,
            "value": req["priority"],
            "requirement": req
        })

    links = [{
        "source": rel["source"],
        "target": rel["target"],
        "name": rel["label"],
        "relationType": rel["type"]
    } for rel in relations]

    categories = [{"name": module} for module in sorted(module_index, key=module_index.get)]
    return {"nodes": nodes, "links": links, "categories": categories}


async def analyze_requirement_document(file: UploadFile) -> Dict[str, Any]:
    parsed = await extract_text_from_upload(file)
    text = parsed["text"]

    agent = create_document_requirement_agent()
    task = f"""
请解析以下需求文档，完成需求识别、功能点拆分、需求类型区分、模块归类、用户故事生成、验收标准生成和需求关系识别。

【文件名】
{parsed['filename']}

【文档正文】
{text[:24000]}
"""

    try:
        result = await agent.run(task=task)
        analysis = _extract_json_object(result.messages[-1].content)
    except Exception:
        analysis = _fallback_requirement_analysis(text)

    normalized = _normalize_analysis(analysis)
    document_id = f"DOC-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
    response = {
        "document_id": document_id,
        "filename": parsed["filename"],
        "extension": parsed["extension"],
        "summary": normalized["summary"],
        "modules": normalized["modules"],
        "requirements": normalized["requirements"],
        "relations": normalized["relations"],
        "graph": _build_graph(normalized["requirements"], normalized["relations"]),
        "statistics": {
            "requirement_count": len(normalized["requirements"]),
            "functional_count": sum(1 for item in normalized["requirements"] if item["type"] == "functional"),
            "non_functional_count": sum(1 for item in normalized["requirements"] if item["type"] == "non_functional"),
            "story_count": sum(len(item.get("user_stories", [])) for item in normalized["requirements"])
        }
    }

    with open(os.path.join(DOCUMENT_STORY_DIR, f"{document_id}.json"), "w", encoding="utf-8") as f:
        json.dump(response, f, ensure_ascii=False, indent=2)

    history_id = save_history("document", {
        **response,
        "title": parsed["filename"],
        "source_document_id": document_id,
    })
    response["history_id"] = history_id

    return response


def load_document_result(document_id: str) -> Dict[str, Any]:
    path = os.path.join(DOCUMENT_STORY_DIR, f"{document_id}.json")
    if not os.path.exists(path):
        raise FileNotFoundError("未找到该文档分析结果")
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_requirement_story(document_id: str, requirement_id: str) -> Dict[str, Any]:
    data = load_document_result(document_id)
    for req in data.get("requirements", []):
        if req.get("id") == requirement_id:
            related = [
                rel for rel in data.get("relations", [])
                if rel.get("source") == requirement_id or rel.get("target") == requirement_id
            ]
            return {"document_id": document_id, "requirement": req, "relations": related}
    raise FileNotFoundError("未找到对应需求")


def export_document_stories_json(document_id: str) -> str:
    data = load_document_result(document_id)
    export_id = f"document_user_stories_{document_id}.json"
    path = os.path.join(DOCUMENT_STORY_DIR, export_id)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return path


def export_document_stories_markdown(document_id: str) -> str:
    data = load_document_result(document_id)
    filename = f"document_user_stories_{document_id}.md"
    path = os.path.join(DOCUMENT_STORY_DIR, filename)
    lines = [
        f"# 需求文档用户故事分析结果：{data.get('filename', '')}",
        "",
        f"## 文档概述\n{data.get('summary', '')}",
        "",
        "## 统计信息",
        f"- 需求总数：{data.get('statistics', {}).get('requirement_count', 0)}",
        f"- 功能性需求：{data.get('statistics', {}).get('functional_count', 0)}",
        f"- 非功能性需求：{data.get('statistics', {}).get('non_functional_count', 0)}",
        f"- 用户故事数：{data.get('statistics', {}).get('story_count', 0)}",
        ""
    ]

    for req in data.get("requirements", []):
        lines.extend([
            f"## {req['id']} {req['title']}",
            f"- 模块：{req['module']}",
            f"- 类型：{'功能性需求' if req['type'] == 'functional' else '非功能性需求'}",
            f"- 优先级：{req['priority']}",
            f"- 描述：{req['description']}",
            ""
        ])
        if req.get("user_stories"):
            lines.append("### 用户故事")
            for story in req["user_stories"]:
                lines.extend([
                    f"#### {story['id']}",
                    story.get("story", ""),
                    "",
                    "**验收标准：**"
                ])
                for item in story.get("acceptance_criteria", []):
                    lines.append(f"- {item}")
                lines.append("")
        if req.get("non_functional_constraints"):
            lines.append("### 非功能性约束")
            for item in req["non_functional_constraints"]:
                lines.append(f"- {item}")
            lines.append("")

    lines.append("## 需求关系")
    for rel in data.get("relations", []):
        lines.append(f"- {rel['source']} --{rel['type']}--> {rel['target']}：{rel.get('label', '')}")

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return path
