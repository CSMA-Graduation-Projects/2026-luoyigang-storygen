import ast
import json
import os
import re
import shutil
import tarfile
import tempfile
import uuid
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

from fastapi import UploadFile

from app.agents.project_story_agent import create_project_story_agent
from app.services.history_service import save_history

PROJECT_STORY_DIR = os.path.join("outputs", "project_story")
os.makedirs(PROJECT_STORY_DIR, exist_ok=True)

SUPPORTED_ARCHIVES = {".zip", ".tar", ".tgz", ".gz"}
SUPPORTED_CODE_EXTENSIONS = {".py", ".js", ".jsx", ".ts", ".tsx", ".vue", ".java"}
IGNORED_DIRS = {
    ".git", ".idea", ".vscode", "node_modules", "dist", "build", "target", "__pycache__",
    ".venv", "venv", "env", ".mvn", ".gradle", "coverage", ".next", ".nuxt"
}
IGNORED_FILES = {"package-lock.json", "yarn.lock", "pnpm-lock.yaml"}
MAX_FILES = 500
MAX_FILE_SIZE = 200_000


class ProjectParseError(ValueError):
    pass


def _safe_filename(filename: str) -> str:
    name = os.path.basename(filename or "project_source.zip")
    return re.sub(r"[^\w.\-\u4e00-\u9fa5]", "_", name)


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk", "gb2312"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _is_safe_path(base: str, target: str) -> bool:
    base_abs = os.path.abspath(base)
    target_abs = os.path.abspath(target)
    return os.path.commonpath([base_abs]) == os.path.commonpath([base_abs, target_abs])


def _extract_zip(data: bytes, workdir: str) -> str:
    archive_path = os.path.join(workdir, "project.zip")
    with open(archive_path, "wb") as f:
        f.write(data)
    extract_dir = os.path.join(workdir, "source")
    os.makedirs(extract_dir, exist_ok=True)

    with zipfile.ZipFile(archive_path) as zf:
        for member in zf.infolist():
            target = os.path.join(extract_dir, member.filename)
            if not _is_safe_path(extract_dir, target):
                raise ProjectParseError("压缩包中存在不安全路径，已拒绝解析")
            zf.extract(member, extract_dir)
    return extract_dir


def _extract_tar(data: bytes, filename: str, workdir: str) -> str:
    archive_path = os.path.join(workdir, filename)
    with open(archive_path, "wb") as f:
        f.write(data)
    extract_dir = os.path.join(workdir, "source")
    os.makedirs(extract_dir, exist_ok=True)

    with tarfile.open(archive_path) as tf:
        for member in tf.getmembers():
            target = os.path.join(extract_dir, member.name)
            if not _is_safe_path(extract_dir, target):
                raise ProjectParseError("压缩包中存在不安全路径，已拒绝解析")
        tf.extractall(extract_dir)
    return extract_dir


async def extract_project_upload(file: UploadFile) -> Dict[str, Any]:
    filename = _safe_filename(file.filename or "project_source.zip")
    ext = Path(filename).suffix.lower()
    data = await file.read()
    if not data:
        raise ProjectParseError("上传项目文件不能为空")

    workdir = tempfile.mkdtemp(prefix="project_story_")
    try:
        if filename.endswith(".tar.gz") or filename.endswith(".tgz"):
            source_dir = _extract_tar(data, filename, workdir)
        elif ext == ".tar":
            source_dir = _extract_tar(data, filename, workdir)
        elif ext == ".zip":
            source_dir = _extract_zip(data, workdir)
        elif ext in SUPPORTED_CODE_EXTENSIONS:
            source_dir = os.path.join(workdir, "source")
            os.makedirs(source_dir, exist_ok=True)
            with open(os.path.join(source_dir, filename), "wb") as f:
                f.write(data)
        else:
            raise ProjectParseError("仅支持上传项目 zip、tar、tar.gz、tgz 压缩包，或单个源码文件")

        return {"filename": filename, "source_dir": source_dir, "workdir": workdir}
    except Exception:
        shutil.rmtree(workdir, ignore_errors=True)
        raise


def _should_skip(path: Path) -> bool:
    if path.name in IGNORED_FILES:
        return True
    return any(part in IGNORED_DIRS for part in path.parts)


def _collect_code_files(source_dir: str) -> List[Path]:
    root = Path(source_dir)
    files: List[Path] = []
    for path in root.rglob("*"):
        if len(files) >= MAX_FILES:
            break
        if not path.is_file() or _should_skip(path):
            continue
        if path.suffix.lower() not in SUPPORTED_CODE_EXTENSIONS:
            continue
        try:
            if path.stat().st_size > MAX_FILE_SIZE:
                continue
        except OSError:
            continue
        files.append(path)
    return files


def _new_node(nodes: List[Dict[str, Any]], **kwargs) -> Dict[str, Any]:
    node_id = f"SYM-{len(nodes) + 1:04d}"
    node = {
        "id": node_id,
        "name": kwargs.get("name", node_id),
        "qualified_name": kwargs.get("qualified_name", kwargs.get("name", node_id)),
        "type": kwargs.get("type", "function"),
        "language": kwargs.get("language", "unknown"),
        "file": kwargs.get("file", ""),
        "line": kwargs.get("line", 1),
        "signature": kwargs.get("signature", ""),
        "code_snippet": kwargs.get("code_snippet", ""),
        "docstring": kwargs.get("docstring", ""),
        "calls": kwargs.get("calls", []),
        "children": kwargs.get("children", []),
        "parent_id": kwargs.get("parent_id"),
        "module": kwargs.get("module", "默认模块"),
        "depth": 0,
    }
    nodes.append(node)
    return node


def _get_source_segment(source: str, node: ast.AST, max_len: int = 1200) -> str:
    try:
        segment = ast.get_source_segment(source, node) or ""
    except Exception:
        segment = ""
    return segment[:max_len]


def _call_name(call: ast.Call) -> Optional[str]:
    func = call.func
    if isinstance(func, ast.Name):
        return func.id
    if isinstance(func, ast.Attribute):
        parts = [func.attr]
        value = func.value
        while isinstance(value, ast.Attribute):
            parts.append(value.attr)
            value = value.value
        if isinstance(value, ast.Name):
            parts.append(value.id)
        return ".".join(reversed(parts))
    return None


class _CallCollector(ast.NodeVisitor):
    def __init__(self):
        self.calls: Set[str] = set()

    def visit_Call(self, node: ast.Call):
        name = _call_name(node)
        if name:
            self.calls.add(name)
        self.generic_visit(node)


def _signature_from_ast(node: Any) -> str:
    if not hasattr(node, "args"):
        return ""
    args = []
    for arg in getattr(node.args, "args", []):
        args.append(arg.arg)
    return f"{node.name}({', '.join(args)})"


def _parse_python_file(path: Path, rel_file: str, text: str, nodes: List[Dict[str, Any]]) -> None:
    try:
        tree = ast.parse(text)
    except SyntaxError:
        return

    def parse_func(fn: ast.AST, parent_name: str = "", parent_id: Optional[str] = None) -> Dict[str, Any]:
        collector = _CallCollector()
        collector.visit(fn)
        name = getattr(fn, "name", "function")
        qualified = f"{parent_name}.{name}" if parent_name else name
        return _new_node(
            nodes,
            name=name,
            qualified_name=qualified,
            type="method" if parent_id else "function",
            language="python",
            file=rel_file,
            line=getattr(fn, "lineno", 1),
            signature=_signature_from_ast(fn),
            code_snippet=_get_source_segment(text, fn),
            docstring=ast.get_docstring(fn) or "",
            calls=sorted(collector.calls),
            parent_id=parent_id,
            module=rel_file.split("/")[0] if "/" in rel_file else Path(rel_file).stem,
        )

    for item in tree.body:
        if isinstance(item, ast.ClassDef):
            class_node = _new_node(
                nodes,
                name=item.name,
                qualified_name=item.name,
                type="class",
                language="python",
                file=rel_file,
                line=getattr(item, "lineno", 1),
                signature=f"class {item.name}",
                code_snippet=_get_source_segment(text, item),
                docstring=ast.get_docstring(item) or "",
                calls=[],
                module=rel_file.split("/")[0] if "/" in rel_file else Path(rel_file).stem,
            )
            for child in item.body:
                if isinstance(child, (ast.FunctionDef, ast.AsyncFunctionDef)):
                    method_node = parse_func(child, item.name, class_node["id"])
                    class_node["children"].append(method_node["id"])
        elif isinstance(item, (ast.FunctionDef, ast.AsyncFunctionDef)):
            parse_func(item)


def _extract_vue_script(text: str) -> str:
    scripts = re.findall(r"<script[^>]*>([\s\S]*?)</script>", text, flags=re.I)
    return "\n".join(scripts) if scripts else text


def _find_calls_regex(body: str) -> List[str]:
    keywords = {"if", "for", "while", "switch", "catch", "function", "return", "typeof", "new", "await", "import", "export"}
    calls = set()
    for match in re.finditer(r"\b([A-Za-z_$][\w$]*)\s*\(", body):
        name = match.group(1)
        if name not in keywords:
            calls.add(name)
    return sorted(calls)


def _parse_js_like_file(path: Path, rel_file: str, text: str, nodes: List[Dict[str, Any]]) -> None:
    language = {".vue": "vue", ".ts": "typescript", ".tsx": "typescript", ".jsx": "javascript"}.get(path.suffix.lower(), "javascript")
    source = _extract_vue_script(text) if path.suffix.lower() == ".vue" else text
    module = rel_file.split("/")[0] if "/" in rel_file else Path(rel_file).stem

    for match in re.finditer(r"class\s+([A-Za-z_$][\w$]*)[\s\S]*?\{", source):
        _new_node(
            nodes,
            name=match.group(1),
            qualified_name=match.group(1),
            type="class",
            language=language,
            file=rel_file,
            line=source[:match.start()].count("\n") + 1,
            signature=f"class {match.group(1)}",
            code_snippet=source[match.start():match.start() + 1000],
            calls=[],
            module=module,
        )

    patterns = [
        r"function\s+([A-Za-z_$][\w$]*)\s*\(([^)]*)\)\s*\{([\s\S]*?)\n\}",
        r"(?:const|let|var)\s+([A-Za-z_$][\w$]*)\s*=\s*(?:async\s*)?\(([^)]*)\)\s*=>\s*\{([\s\S]*?)\n\}",
        r"(?:const|let|var)\s+([A-Za-z_$][\w$]*)\s*=\s*(?:async\s*)?function\s*\(([^)]*)\)\s*\{([\s\S]*?)\n\}",
    ]
    seen: Set[Tuple[str, int]] = set()
    for pattern in patterns:
        for match in re.finditer(pattern, source):
            name = match.group(1)
            key = (name, match.start())
            if key in seen:
                continue
            seen.add(key)
            args = match.group(2).strip()
            body = match.group(3)
            _new_node(
                nodes,
                name=name,
                qualified_name=name,
                type="function",
                language=language,
                file=rel_file,
                line=source[:match.start()].count("\n") + 1,
                signature=f"{name}({args})",
                code_snippet=source[match.start():match.end()][:1200],
                calls=_find_calls_regex(body),
                module=module,
            )


def _parse_java_file(path: Path, rel_file: str, text: str, nodes: List[Dict[str, Any]]) -> None:
    module = rel_file.split("/")[0] if "/" in rel_file else Path(rel_file).stem
    classes: List[Dict[str, Any]] = []
    for match in re.finditer(r"\bclass\s+([A-Za-z_][\w]*)", text):
        node = _new_node(
            nodes,
            name=match.group(1),
            qualified_name=match.group(1),
            type="class",
            language="java",
            file=rel_file,
            line=text[:match.start()].count("\n") + 1,
            signature=f"class {match.group(1)}",
            code_snippet=text[match.start():match.start() + 1200],
            module=module,
        )
        classes.append(node)

    method_pattern = r"(?:public|private|protected|static|final|synchronized|abstract|\s)+[\w<>\[\]]+\s+([A-Za-z_][\w]*)\s*\(([^)]*)\)\s*\{([\s\S]*?)\n\s*\}"
    for match in re.finditer(method_pattern, text):
        name = match.group(1)
        if name in {"if", "for", "while", "switch", "catch"}:
            continue
        parent_id = classes[0]["id"] if classes else None
        node = _new_node(
            nodes,
            name=name,
            qualified_name=f"{classes[0]['name']}.{name}" if classes else name,
            type="method" if parent_id else "function",
            language="java",
            file=rel_file,
            line=text[:match.start()].count("\n") + 1,
            signature=f"{name}({match.group(2).strip()})",
            code_snippet=text[match.start():match.end()][:1200],
            calls=_find_calls_regex(match.group(3)),
            parent_id=parent_id,
            module=module,
        )
        if classes:
            classes[0]["children"].append(node["id"])


def parse_project_symbols(source_dir: str) -> Dict[str, Any]:
    root = Path(source_dir)
    code_files = _collect_code_files(source_dir)
    nodes: List[Dict[str, Any]] = []
    file_summaries = []

    for path in code_files:
        rel_file = path.relative_to(root).as_posix()
        try:
            text = _decode_text(path.read_bytes())
        except Exception:
            continue
        before_count = len(nodes)
        suffix = path.suffix.lower()
        if suffix == ".py":
            _parse_python_file(path, rel_file, text, nodes)
        elif suffix in {".js", ".jsx", ".ts", ".tsx", ".vue"}:
            _parse_js_like_file(path, rel_file, text, nodes)
        elif suffix == ".java":
            _parse_java_file(path, rel_file, text, nodes)
        after_count = len(nodes)
        if after_count > before_count:
            file_summaries.append({
                "file": rel_file,
                "language": suffix.replace(".", ""),
                "symbol_count": after_count - before_count,
                "module": rel_file.split("/")[0] if "/" in rel_file else Path(rel_file).stem,
            })

    if not nodes:
        raise ProjectParseError("未在项目中识别到可分析的函数或类，请确认上传的是源码项目")

    return {"nodes": nodes, "files": file_summaries}


def _build_call_edges(nodes: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    by_simple: Dict[str, List[Dict[str, Any]]] = {}
    by_qualified: Dict[str, Dict[str, Any]] = {}
    for node in nodes:
        by_simple.setdefault(node["name"], []).append(node)
        by_qualified[node["qualified_name"]] = node

    edges = []
    seen = set()
    for source in nodes:
        for call in source.get("calls", []):
            simple = call.split(".")[-1]
            candidates = []
            if call in by_qualified:
                candidates = [by_qualified[call]]
            elif simple in by_simple:
                same_file = [item for item in by_simple[simple] if item["file"] == source["file"]]
                candidates = same_file or by_simple[simple]
            if not candidates:
                continue
            target = candidates[0]
            if target["id"] == source["id"]:
                continue
            key = (source["id"], target["id"])
            if key in seen:
                continue
            seen.add(key)
            edges.append({
                "source": source["id"],
                "target": target["id"],
                "type": "calls",
                "label": f"{source['name']} 调用 {target['name']}"
            })
    return edges


def _compute_depth(nodes: List[Dict[str, Any]], edges: List[Dict[str, Any]]) -> None:
    outgoing: Dict[str, List[str]] = {}
    for edge in edges:
        outgoing.setdefault(edge["source"], []).append(edge["target"])

    memo: Dict[str, int] = {}

    def depth(node_id: str, visiting: Set[str]) -> int:
        if node_id in memo:
            return memo[node_id]
        if node_id in visiting:
            return 0
        visiting.add(node_id)
        children = outgoing.get(node_id, [])
        if not children:
            memo[node_id] = 0
        else:
            memo[node_id] = 1 + max(depth(child, visiting) for child in children)
        visiting.remove(node_id)
        return memo[node_id]

    for node in nodes:
        node["depth"] = depth(node["id"], set())


def _build_function_tree(nodes: List[Dict[str, Any]], edges: List[Dict[str, Any]], files: List[Dict[str, Any]]) -> Dict[str, Any]:
    node_map = {node["id"]: node for node in nodes}
    targets = {edge["target"] for edge in edges}
    roots = [node for node in nodes if node["id"] not in targets]
    if not roots:
        roots = sorted(nodes, key=lambda item: item.get("depth", 0), reverse=True)[:10]

    outgoing: Dict[str, List[str]] = {}
    for edge in edges:
        outgoing.setdefault(edge["source"], []).append(edge["target"])

    def to_tree(node: Dict[str, Any], visiting: Set[str]) -> Dict[str, Any]:
        tree = {
            "id": node["id"],
            "name": node["name"],
            "label": f"{node['name']}｜{node['type']}",
            "type": node["type"],
            "file": node["file"],
            "line": node["line"],
            "depth": node.get("depth", 0),
            "node": node,
            "children": []
        }
        if node["id"] in visiting:
            tree["children"].append({"id": f"cycle-{node['id']}", "name": "循环调用已省略", "type": "cycle"})
            return tree
        visiting.add(node["id"])
        child_ids = outgoing.get(node["id"], [])
        child_ids = sorted(child_ids, key=lambda cid: node_map[cid].get("depth", 0), reverse=True)
        for child_id in child_ids:
            tree["children"].append(to_tree(node_map[child_id], visiting))
        visiting.remove(node["id"])
        return tree

    return {
        "id": "PROJECT_ROOT",
        "name": "项目调用树",
        "type": "project",
        "children": [to_tree(node, set()) for node in sorted(roots, key=lambda item: item.get("depth", 0), reverse=True)[:30]],
        "files": files,
    }


def _extract_json_object(text: str) -> Dict[str, Any]:
    cleaned = text.strip()
    cleaned = re.sub(r"^```json\s*", "", cleaned)
    cleaned = re.sub(r"^```\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", cleaned)
        if not match:
            raise
        return json.loads(match.group(0))


def _fallback_project_stories(nodes: List[Dict[str, Any]]) -> Dict[str, Any]:
    sorted_nodes = sorted(nodes, key=lambda item: (item.get("depth", 0), item.get("file", "")))
    selected = sorted_nodes[:80]
    stories = []
    for index, node in enumerate(selected, start=1):
        is_bottom = node.get("depth", 0) == 0
        capability = "提供底层能力支撑" if is_bottom else "组合下层能力完成业务处理"
        role = "系统用户" if node["type"] in {"class", "function", "method"} else "开发人员"
        want = f"通过 {node['name']} {capability}"
        benefit = "保证系统功能能够被正确执行并形成可验证的业务流程"
        stories.append({
            "id": f"US-{index:03d}",
            "node_id": node["id"],
            "node_name": node["name"],
            "level": node["type"],
            "module": node.get("module", "默认模块"),
            "role": role,
            "want": want,
            "benefit": benefit,
            "story": f"As a {role}, I want {want}, So that {benefit}",
            "acceptance_criteria": [
                f"Given 项目运行环境正常，When 调用 {node['name']} 对应功能，Then 系统应按照函数职责返回正确结果。",
                f"Given 输入数据异常，When 执行 {node['name']}，Then 系统应给出可控处理结果，避免流程中断。"
            ],
            "technical_reasoning": f"该故事由 {node.get('file')} 第 {node.get('line')} 行的 {node.get('type')} 节点自底向上归纳得到。"
        })
    modules = sorted({story["module"] for story in stories})
    return {
        "summary": "系统已根据源码函数、类与调用关系，自底向上生成用户故事。",
        "modules": modules,
        "stories": stories
    }


def _normalize_project_stories(raw: Dict[str, Any], nodes: List[Dict[str, Any]]) -> Dict[str, Any]:
    valid_ids = {node["id"] for node in nodes}
    node_by_id = {node["id"]: node for node in nodes}
    stories = []
    for index, item in enumerate(raw.get("stories") or [], start=1):
        node_id = item.get("node_id")
        if node_id not in valid_ids:
            continue
        node = node_by_id[node_id]
        criteria = item.get("acceptance_criteria") or []
        if isinstance(criteria, str):
            criteria = [criteria]
        stories.append({
            "id": item.get("id") or f"US-{index:03d}",
            "node_id": node_id,
            "node_name": item.get("node_name") or node["name"],
            "level": item.get("level") or node["type"],
            "module": item.get("module") or node.get("module", "默认模块"),
            "role": item.get("role") or "系统用户",
            "want": item.get("want") or f"使用 {node['name']} 完成对应功能",
            "benefit": item.get("benefit") or "满足业务处理目标",
            "story": item.get("story") or f"As a 系统用户, I want 使用 {node['name']} 完成对应功能, So that 满足业务处理目标",
            "acceptance_criteria": criteria,
            "technical_reasoning": item.get("technical_reasoning") or "由源码调用树分析得到。"
        })
    if not stories:
        return _fallback_project_stories(nodes)
    return {
        "summary": raw.get("summary") or "已完成项目源码解析与用户故事生成。",
        "modules": raw.get("modules") or sorted({story["module"] for story in stories}),
        "stories": stories
    }


def _build_story_tree(function_tree: Dict[str, Any], stories: List[Dict[str, Any]]) -> Dict[str, Any]:
    story_by_node: Dict[str, List[Dict[str, Any]]] = {}
    for story in stories:
        story_by_node.setdefault(story["node_id"], []).append(story)

    def convert(tree_node: Dict[str, Any]) -> Dict[str, Any]:
        node_id = tree_node.get("id")
        attached = story_by_node.get(node_id, [])
        label = tree_node.get("name", "节点")
        if attached:
            label = f"{label}\n{attached[0]['id']}"
        return {
            "id": f"STORY-{node_id}",
            "node_id": node_id,
            "name": label,
            "type": "story_node",
            "stories": attached,
            "children": [convert(child) for child in tree_node.get("children", []) if child.get("type") != "cycle"]
        }

    return convert(function_tree)


def _build_graph(nodes: List[Dict[str, Any]], edges: List[Dict[str, Any]]) -> Dict[str, Any]:
    categories = sorted({node["type"] for node in nodes})
    category_index = {name: index for index, name in enumerate(categories)}
    return {
        "nodes": [
            {
                "id": node["id"],
                "name": f"{node['name']}\n{node['type']}",
                "category": category_index.get(node["type"], 0),
                "symbolSize": 58 if node["type"] == "class" else 44,
                "value": node.get("depth", 0),
                "node": node,
            }
            for node in nodes
        ],
        "links": [
            {"source": edge["source"], "target": edge["target"], "name": edge["label"], "type": edge["type"]}
            for edge in edges
        ],
        "categories": [{"name": name} for name in categories],
    }


async def analyze_project_source(file: UploadFile) -> Dict[str, Any]:
    extracted = await extract_project_upload(file)
    try:
        parsed = parse_project_symbols(extracted["source_dir"])
        nodes = parsed["nodes"]
        edges = _build_call_edges(nodes)
        _compute_depth(nodes, edges)
        function_tree = _build_function_tree(nodes, edges, parsed["files"])

        compact_nodes = [
            {
                "id": node["id"],
                "name": node["name"],
                "qualified_name": node["qualified_name"],
                "type": node["type"],
                "language": node["language"],
                "file": node["file"],
                "line": node["line"],
                "signature": node["signature"],
                "docstring": node.get("docstring", ""),
                "calls": node.get("calls", [])[:20],
                "depth": node.get("depth", 0),
                "module": node.get("module", "默认模块"),
            }
            for node in sorted(nodes, key=lambda item: item.get("depth", 0))[:120]
        ]
        compact_edges = edges[:200]
        agent = create_project_story_agent()
        task = f"""
请根据以下项目源码分析结果，自底向上生成用户故事。

【项目文件】
{extracted['filename']}

【函数/类节点，已按底层到上层排序】
{json.dumps(compact_nodes, ensure_ascii=False, indent=2)}

【调用关系】
{json.dumps(compact_edges, ensure_ascii=False, indent=2)}
"""
        try:
            result = await agent.run(task=task)
            raw_stories = _extract_json_object(result.messages[-1].content)
            story_result = _normalize_project_stories(raw_stories, nodes)
        except Exception:
            story_result = _fallback_project_stories(nodes)

        project_id = f"PROJ-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
        response = {
            "project_id": project_id,
            "filename": extracted["filename"],
            "summary": story_result["summary"],
            "modules": story_result["modules"],
            "files": parsed["files"],
            "nodes": nodes,
            "edges": edges,
            "function_tree": function_tree,
            "stories": story_result["stories"],
            "story_tree": _build_story_tree(function_tree, story_result["stories"]),
            "graph": _build_graph(nodes, edges),
            "statistics": {
                "file_count": len(parsed["files"]),
                "symbol_count": len(nodes),
                "class_count": sum(1 for node in nodes if node["type"] == "class"),
                "function_count": sum(1 for node in nodes if node["type"] in {"function", "method"}),
                "edge_count": len(edges),
                "story_count": len(story_result["stories"]),
            }
        }

        with open(os.path.join(PROJECT_STORY_DIR, f"{project_id}.json"), "w", encoding="utf-8") as f:
            json.dump(response, f, ensure_ascii=False, indent=2)

        history_id = save_history("project", {
            **response,
            "title": extracted["filename"],
            "source_project_id": project_id,
        })
        response["history_id"] = history_id
        return response
    finally:
        shutil.rmtree(extracted.get("workdir", ""), ignore_errors=True)


def load_project_result(project_id: str) -> Dict[str, Any]:
    path = os.path.join(PROJECT_STORY_DIR, f"{project_id}.json")
    if not os.path.exists(path):
        raise FileNotFoundError("未找到该项目源码分析结果")
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_project_node_story(project_id: str, node_id: str) -> Dict[str, Any]:
    data = load_project_result(project_id)
    node = next((item for item in data.get("nodes", []) if item.get("id") == node_id), None)
    if not node:
        raise FileNotFoundError("未找到对应函数或类节点")
    stories = [story for story in data.get("stories", []) if story.get("node_id") == node_id]
    incoming = [edge for edge in data.get("edges", []) if edge.get("target") == node_id]
    outgoing = [edge for edge in data.get("edges", []) if edge.get("source") == node_id]
    return {"project_id": project_id, "node": node, "stories": stories, "incoming": incoming, "outgoing": outgoing}


def export_project_stories_json(project_id: str) -> str:
    data = load_project_result(project_id)
    filename = f"project_user_stories_{project_id}.json"
    path = os.path.join(PROJECT_STORY_DIR, filename)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return path


def export_project_stories_markdown(project_id: str) -> str:
    data = load_project_result(project_id)
    filename = f"project_user_stories_{project_id}.md"
    path = os.path.join(PROJECT_STORY_DIR, filename)
    node_by_id = {node["id"]: node for node in data.get("nodes", [])}
    lines = [
        f"# 项目源码用户故事分析结果：{data.get('filename', '')}",
        "",
        f"## 项目概述\n{data.get('summary', '')}",
        "",
        "## 统计信息",
        f"- 源码文件数：{data.get('statistics', {}).get('file_count', 0)}",
        f"- 函数/类节点数：{data.get('statistics', {}).get('symbol_count', 0)}",
        f"- 调用关系数：{data.get('statistics', {}).get('edge_count', 0)}",
        f"- 用户故事数：{data.get('statistics', {}).get('story_count', 0)}",
        "",
        "## 用户故事",
    ]
    for story in data.get("stories", []):
        node = node_by_id.get(story.get("node_id"), {})
        lines.extend([
            f"### {story['id']} {story.get('node_name', '')}",
            f"- 源码节点：{story.get('node_id')} / {node.get('type', '')}",
            f"- 文件位置：{node.get('file', '')}:{node.get('line', '')}",
            f"- 模块：{story.get('module', '')}",
            f"- 用户故事：{story.get('story', '')}",
            f"- 技术归因：{story.get('technical_reasoning', '')}",
            "",
            "**验收标准：**",
        ])
        for item in story.get("acceptance_criteria", []):
            lines.append(f"- {item}")
        lines.append("")

    lines.append("## 调用关系")
    for edge in data.get("edges", []):
        source = node_by_id.get(edge.get("source"), {}).get("qualified_name", edge.get("source"))
        target = node_by_id.get(edge.get("target"), {}).get("qualified_name", edge.get("target"))
        lines.append(f"- {source} -> {target}")

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return path

# =========================
# 逐节点 AI 生成增强版实现
# =========================
from html import escape as _html_escape


def _story_fallback_for_node(node: Dict[str, Any], index: int) -> Dict[str, Any]:
    is_bottom = node.get("depth", 0) == 0
    capability = "提供底层能力支撑" if is_bottom else "组合下层节点完成业务处理"
    role = "系统用户"
    want = f"使用 {node.get('qualified_name') or node.get('name')} {capability}"
    benefit = "让系统功能能够按源码中的业务逻辑稳定、可验证地运行"
    return {
        "id": f"US-{index:03d}",
        "node_id": node["id"],
        "node_name": node.get("name", ""),
        "level": node.get("type", "function"),
        "module": node.get("module", "默认模块"),
        "role": role,
        "want": want,
        "benefit": benefit,
        "story": f"As a {role}, I want {want}, So that {benefit}",
        "acceptance_criteria": [
            f"Given 系统运行环境正常，When 触发 {node.get('name')} 对应功能，Then 系统应按该节点职责完成处理并返回可验证结果。",
            f"Given 输入或依赖数据异常，When 执行 {node.get('name')}，Then 系统应进行可控处理，避免项目流程异常中断。"
        ],
        "technical_reasoning": f"该用户故事由 {node.get('file')} 第 {node.get('line')} 行的 {node.get('type')} 节点、源码片段和调用关系逆向归纳得到。",
        "ai_generated": False,
    }


def _normalize_single_node_story(raw: Dict[str, Any], node: Dict[str, Any], index: int) -> Dict[str, Any]:
    source = raw
    if isinstance(raw.get("stories"), list) and raw["stories"]:
        source = raw["stories"][0]
    if not isinstance(source, dict):
        return _story_fallback_for_node(node, index)
    criteria = source.get("acceptance_criteria") or []
    if isinstance(criteria, str):
        criteria = [criteria]
    if not criteria:
        criteria = _story_fallback_for_node(node, index)["acceptance_criteria"]
    story = {
        "id": source.get("id") or f"US-{index:03d}",
        "node_id": node["id"],
        "node_name": source.get("node_name") or node.get("name", ""),
        "level": source.get("level") or node.get("type", "function"),
        "module": source.get("module") or node.get("module", "默认模块"),
        "role": source.get("role") or "系统用户",
        "want": source.get("want") or f"使用 {node.get('qualified_name') or node.get('name')} 完成对应功能",
        "benefit": source.get("benefit") or "满足项目源码体现的业务处理目标",
        "story": source.get("story") or f"As a 系统用户, I want 使用 {node.get('name')} 完成对应功能, So that 满足项目源码体现的业务处理目标",
        "acceptance_criteria": criteria,
        "technical_reasoning": source.get("technical_reasoning") or "由源码节点职责、源码片段和调用关系逆向分析得到。",
        "ai_generated": True,
    }
    return story


def _relations_for_node(node: Dict[str, Any], nodes: List[Dict[str, Any]], edges: List[Dict[str, Any]]) -> Dict[str, Any]:
    node_map = {item["id"]: item for item in nodes}
    incoming = []
    outgoing = []
    for edge in edges:
        if edge.get("target") == node["id"]:
            src = node_map.get(edge.get("source"), {})
            incoming.append({"id": src.get("id"), "name": src.get("qualified_name") or src.get("name"), "file": src.get("file")})
        if edge.get("source") == node["id"]:
            tgt = node_map.get(edge.get("target"), {})
            outgoing.append({"id": tgt.get("id"), "name": tgt.get("qualified_name") or tgt.get("name"), "file": tgt.get("file")})
    return {"incoming": incoming, "outgoing": outgoing}


async def _generate_story_for_node_with_ai(node: Dict[str, Any], nodes: List[Dict[str, Any]], edges: List[Dict[str, Any]], index: int) -> Dict[str, Any]:
    relations = _relations_for_node(node, nodes, edges)
    task = f"""
请只为下面这一个项目源码节点逆向生成 1 条用户故事。必须使用 AI 根据源码职责和调用关系生成，不要只复述函数名。

要求：
1. 必须输出合法 JSON，不要输出 Markdown，不要输出解释文字。
2. stories 数组中必须且只能包含 1 条用户故事。
3. node_id 必须等于 {node['id']}。
4. 用户故事要体现该节点在项目中的功能价值；底层函数可以写成支撑型用户故事。
5. acceptance_criteria 使用 Given-When-Then，至少 2 条。

【当前节点】
{json.dumps({k: node.get(k) for k in ['id','name','qualified_name','type','language','file','line','signature','docstring','calls','depth','module']}, ensure_ascii=False, indent=2)}

【源码片段】
{node.get('code_snippet', '')[:1600]}

【调用下层节点】
{json.dumps(relations['outgoing'], ensure_ascii=False, indent=2)}

【被上层节点调用】
{json.dumps(relations['incoming'], ensure_ascii=False, indent=2)}

请按此 JSON 格式返回：
{{
  "summary": "当前源码节点功能概述",
  "modules": ["{node.get('module', '默认模块')}"],
  "stories": [
    {{
      "id": "US-{index:03d}",
      "node_id": "{node['id']}",
      "node_name": "{node.get('name','')}",
      "level": "{node.get('type','function')}",
      "module": "{node.get('module','默认模块')}",
      "role": "用户角色",
      "want": "希望完成的功能",
      "benefit": "业务价值",
      "story": "As a <角色>, I want <功能>, So that <价值>",
      "acceptance_criteria": ["Given ... When ... Then ...", "Given ... When ... Then ..."],
      "technical_reasoning": "说明该故事如何从源码职责和调用关系推导得到"
    }}
  ]
}}
"""
    try:
        agent = create_project_story_agent()
        result = await agent.run(task=task)
        raw = _extract_json_object(result.messages[-1].content)
        return _normalize_single_node_story(raw, node, index)
    except Exception as exc:
        story = _story_fallback_for_node(node, index)
        story["technical_reasoning"] += f"（AI 生成异常，已使用兜底规则：{str(exc)[:120]}）"
        story["ai_error"] = str(exc)[:300]
        return story


def _build_project_response(project_id: str, filename: str, parsed: Dict[str, Any], nodes: List[Dict[str, Any]], edges: List[Dict[str, Any]], function_tree: Dict[str, Any], stories: List[Dict[str, Any]]) -> Dict[str, Any]:
    modules = sorted({story.get("module", "默认模块") for story in stories})
    ai_count = sum(1 for story in stories if story.get("ai_generated"))
    response = {
        "project_id": project_id,
        "filename": filename,
        "summary": f"系统已按项目源码节点逐一逆向生成用户故事，共识别 {len(nodes)} 个节点，已生成 {len(stories)} 条用户故事。",
        "modules": modules,
        "files": parsed["files"],
        "nodes": nodes,
        "edges": edges,
        "function_tree": function_tree,
        "stories": stories,
        "story_tree": _build_story_tree(function_tree, stories),
        "graph": _build_graph(nodes, edges),
        "statistics": {
            "file_count": len(parsed["files"]),
            "symbol_count": len(nodes),
            "class_count": sum(1 for node in nodes if node["type"] == "class"),
            "function_count": sum(1 for node in nodes if node["type"] in {"function", "method"}),
            "edge_count": len(edges),
            "story_count": len(stories),
            "ai_story_count": ai_count,
            "fallback_story_count": len(stories) - ai_count,
        }
    }
    return response


def _save_project_response(response: Dict[str, Any]) -> Dict[str, Any]:
    project_id = response["project_id"]
    with open(os.path.join(PROJECT_STORY_DIR, f"{project_id}.json"), "w", encoding="utf-8") as f:
        json.dump(response, f, ensure_ascii=False, indent=2)
    history_id = save_history("project", {
        **response,
        "title": response.get("filename", "项目源码"),
        "source_project_id": project_id,
    })
    response["history_id"] = history_id
    with open(os.path.join(PROJECT_STORY_DIR, f"{project_id}.json"), "w", encoding="utf-8") as f:
        json.dump(response, f, ensure_ascii=False, indent=2)
    return response


async def analyze_project_source(file: UploadFile) -> Dict[str, Any]:
    """非流式接口：仍然逐节点调用 AI，确保每个源码节点都有用户故事。"""
    extracted = await extract_project_upload(file)
    try:
        parsed = parse_project_symbols(extracted["source_dir"])
        nodes = parsed["nodes"]
        edges = _build_call_edges(nodes)
        _compute_depth(nodes, edges)
        function_tree = _build_function_tree(nodes, edges, parsed["files"])
        ordered_nodes = sorted(nodes, key=lambda item: (item.get("depth", 0), item.get("file", ""), item.get("line", 0)))
        stories = []
        for index, node in enumerate(ordered_nodes, start=1):
            stories.append(await _generate_story_for_node_with_ai(node, nodes, edges, index))
        project_id = f"PROJ-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
        response = _build_project_response(project_id, extracted["filename"], parsed, nodes, edges, function_tree, stories)
        return _save_project_response(response)
    finally:
        shutil.rmtree(extracted.get("workdir", ""), ignore_errors=True)


def _sse_event(event: str, data: Dict[str, Any]) -> str:
    payload = {"event": event, **data}
    return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"


async def stream_project_source(file: UploadFile):
    """流式接口：解析完成后按节点逐条返回 AI 生成结果。"""
    extracted = None
    try:
        extracted = await extract_project_upload(file)
        yield _sse_event("start", {"filename": extracted["filename"], "message": "开始解析项目源码"})
        parsed = parse_project_symbols(extracted["source_dir"])
        nodes = parsed["nodes"]
        edges = _build_call_edges(nodes)
        _compute_depth(nodes, edges)
        function_tree = _build_function_tree(nodes, edges, parsed["files"])
        project_id = f"PROJ-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
        yield _sse_event("parsed", {
            "project_id": project_id,
            "filename": extracted["filename"],
            "files": parsed["files"],
            "nodes": nodes,
            "edges": edges,
            "function_tree": function_tree,
            "graph": _build_graph(nodes, edges),
            "statistics": {
                "file_count": len(parsed["files"]),
                "symbol_count": len(nodes),
                "class_count": sum(1 for node in nodes if node["type"] == "class"),
                "function_count": sum(1 for node in nodes if node["type"] in {"function", "method"}),
                "edge_count": len(edges),
                "story_count": 0,
                "ai_story_count": 0,
                "fallback_story_count": 0,
            }
        })
        ordered_nodes = sorted(nodes, key=lambda item: (item.get("depth", 0), item.get("file", ""), item.get("line", 0)))
        stories: List[Dict[str, Any]] = []
        total = len(ordered_nodes)
        for index, node in enumerate(ordered_nodes, start=1):
            yield _sse_event("node_start", {"index": index, "total": total, "node": node})
            story = await _generate_story_for_node_with_ai(node, nodes, edges, index)
            stories.append(story)
            yield _sse_event("node_story", {
                "index": index,
                "total": total,
                "node_id": node["id"],
                "story": story,
                "progress": round(index * 100 / total, 2),
            })
        response = _build_project_response(project_id, extracted["filename"], parsed, nodes, edges, function_tree, stories)
        response = _save_project_response(response)
        yield _sse_event("done", {"result": response})
    except Exception as exc:
        yield _sse_event("error", {"message": f"项目源码分析失败：{exc}"})
    finally:
        if extracted:
            shutil.rmtree(extracted.get("workdir", ""), ignore_errors=True)


def _project_markdown_content(data: Dict[str, Any]) -> str:
    node_by_id = {node["id"]: node for node in data.get("nodes", [])}
    lines = [
        f"# 项目源码用户故事分析结果：{data.get('filename', '')}",
        "",
        f"## 一、项目概述\n{data.get('summary', '')}",
        "",
        "## 二、统计信息",
        f"- 源码文件数：{data.get('statistics', {}).get('file_count', 0)}",
        f"- 函数/类节点数：{data.get('statistics', {}).get('symbol_count', 0)}",
        f"- 调用关系数：{data.get('statistics', {}).get('edge_count', 0)}",
        f"- 用户故事数：{data.get('statistics', {}).get('story_count', 0)}",
        f"- AI 生成故事数：{data.get('statistics', {}).get('ai_story_count', 0)}",
        "",
        "## 三、逐节点用户故事",
    ]
    for story in data.get("stories", []):
        node = node_by_id.get(story.get("node_id"), {})
        lines.extend([
            f"### {story.get('id')} {story.get('node_name', '')}",
            f"- 源码节点：{story.get('node_id')} / {node.get('type', '')}",
            f"- 文件位置：{node.get('file', '')}:{node.get('line', '')}",
            f"- 函数签名：`{node.get('signature', '')}`",
            f"- 模块：{story.get('module', '')}",
            f"- 是否 AI 生成：{'是' if story.get('ai_generated') else '否，使用兜底规则'}",
            f"- 用户故事：{story.get('story', '')}",
            f"- 技术归因：{story.get('technical_reasoning', '')}",
            "",
            "**验收标准：**",
        ])
        for item in story.get("acceptance_criteria", []):
            lines.append(f"- {item}")
        lines.append("")
    lines.append("## 四、调用关系")
    if data.get("edges"):
        for edge in data.get("edges", []):
            source = node_by_id.get(edge.get("source"), {}).get("qualified_name", edge.get("source"))
            target = node_by_id.get(edge.get("target"), {}).get("qualified_name", edge.get("target"))
            lines.append(f"- {source} -> {target}")
    else:
        lines.append("- 未识别到函数之间的调用关系。")
    return "\n".join(lines)


def export_project_stories_markdown(project_id: str) -> str:
    data = load_project_result(project_id)
    filename = f"project_user_stories_{project_id}.md"
    path = os.path.join(PROJECT_STORY_DIR, filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(_project_markdown_content(data))
    return path


def export_project_stories_docx(project_id: str) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise ProjectParseError("导出 Word 需要安装 python-docx：pip install python-docx") from exc
    data = load_project_result(project_id)
    node_by_id = {node["id"]: node for node in data.get("nodes", [])}
    filename = f"project_user_stories_{project_id}.docx"
    path = os.path.join(PROJECT_STORY_DIR, filename)
    doc = Document()
    doc.add_heading(f"项目源码用户故事分析结果：{data.get('filename', '')}", 0)
    doc.add_heading("一、项目概述", level=1)
    doc.add_paragraph(data.get("summary", ""))
    doc.add_heading("二、统计信息", level=1)
    stats = data.get("statistics", {})
    for line in [
        f"源码文件数：{stats.get('file_count', 0)}",
        f"函数/类节点数：{stats.get('symbol_count', 0)}",
        f"调用关系数：{stats.get('edge_count', 0)}",
        f"用户故事数：{stats.get('story_count', 0)}",
        f"AI 生成故事数：{stats.get('ai_story_count', 0)}",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("三、逐节点用户故事", level=1)
    for story in data.get("stories", []):
        node = node_by_id.get(story.get("node_id"), {})
        doc.add_heading(f"{story.get('id')} {story.get('node_name', '')}", level=2)
        doc.add_paragraph(f"源码节点：{story.get('node_id')} / {node.get('type', '')}")
        doc.add_paragraph(f"文件位置：{node.get('file', '')}:{node.get('line', '')}")
        doc.add_paragraph(f"函数签名：{node.get('signature', '')}")
        doc.add_paragraph(f"模块：{story.get('module', '')}")
        doc.add_paragraph(f"用户故事：{story.get('story', '')}")
        doc.add_paragraph(f"技术归因：{story.get('technical_reasoning', '')}")
        doc.add_paragraph("验收标准：")
        for item in story.get("acceptance_criteria", []):
            doc.add_paragraph(item, style="List Number")
    doc.add_heading("四、调用关系", level=1)
    if data.get("edges"):
        for edge in data.get("edges", []):
            source = node_by_id.get(edge.get("source"), {}).get("qualified_name", edge.get("source"))
            target = node_by_id.get(edge.get("target"), {}).get("qualified_name", edge.get("target"))
            doc.add_paragraph(f"{source} -> {target}", style="List Bullet")
    else:
        doc.add_paragraph("未识别到函数之间的调用关系。")
    doc.save(path)
    return path


def export_project_stories_pdf(project_id: str) -> str:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase import pdfmetrics
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except Exception as exc:
        raise ProjectParseError("导出 PDF 需要安装 reportlab：pip install reportlab") from exc
    data = load_project_result(project_id)
    node_by_id = {node["id"]: node for node in data.get("nodes", [])}
    filename = f"project_user_stories_{project_id}.pdf"
    path = os.path.join(PROJECT_STORY_DIR, filename)
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title = ParagraphStyle("TitleCN", parent=styles["Title"], fontName="STSong-Light", fontSize=18, leading=24)
    heading = ParagraphStyle("HeadingCN", parent=styles["Heading2"], fontName="STSong-Light", fontSize=14, leading=20, spaceBefore=10)
    body = ParagraphStyle("BodyCN", parent=styles["BodyText"], fontName="STSong-Light", fontSize=10.5, leading=17)

    def p(text: Any, style=body):
        return Paragraph(_html_escape(str(text or "")).replace("\n", "<br/>"), style)

    content = [p(f"项目源码用户故事分析结果：{data.get('filename', '')}", title), Spacer(1, 12)]
    content.extend([p("一、项目概述", heading), p(data.get("summary", "")), Spacer(1, 8)])
    stats = data.get("statistics", {})
    content.extend([
        p("二、统计信息", heading),
        p(f"源码文件数：{stats.get('file_count', 0)}"),
        p(f"函数/类节点数：{stats.get('symbol_count', 0)}"),
        p(f"调用关系数：{stats.get('edge_count', 0)}"),
        p(f"用户故事数：{stats.get('story_count', 0)}"),
        p(f"AI 生成故事数：{stats.get('ai_story_count', 0)}"),
        Spacer(1, 8),
        p("三、逐节点用户故事", heading),
    ])
    for story in data.get("stories", []):
        node = node_by_id.get(story.get("node_id"), {})
        content.extend([
            p(f"{story.get('id')} {story.get('node_name', '')}", heading),
            p(f"源码节点：{story.get('node_id')} / {node.get('type', '')}"),
            p(f"文件位置：{node.get('file', '')}:{node.get('line', '')}"),
            p(f"函数签名：{node.get('signature', '')}"),
            p(f"模块：{story.get('module', '')}"),
            p(f"用户故事：{story.get('story', '')}"),
            p(f"技术归因：{story.get('technical_reasoning', '')}"),
            p("验收标准："),
        ])
        for item in story.get("acceptance_criteria", []):
            content.append(p(f"• {item}"))
        content.append(Spacer(1, 8))
    content.append(p("四、调用关系", heading))
    if data.get("edges"):
        for edge in data.get("edges", []):
            source = node_by_id.get(edge.get("source"), {}).get("qualified_name", edge.get("source"))
            target = node_by_id.get(edge.get("target"), {}).get("qualified_name", edge.get("target"))
            content.append(p(f"• {source} -> {target}"))
    else:
        content.append(p("未识别到函数之间的调用关系。"))
    SimpleDocTemplate(path, pagesize=A4).build(content)
    return path


def export_project_story(project_id: str, format: str = "markdown") -> str:
    safe_format = (format or "markdown").lower()
    if safe_format in {"md", "markdown"}:
        return export_project_stories_markdown(project_id)
    if safe_format == "docx":
        return export_project_stories_docx(project_id)
    if safe_format == "pdf":
        return export_project_stories_pdf(project_id)
    if safe_format == "json":
        return export_project_stories_json(project_id)
    raise ProjectParseError("导出格式仅支持 markdown、docx、pdf、json")
